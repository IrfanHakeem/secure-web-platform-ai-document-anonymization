import csv
import io
import uuid
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document as DocxDocument
from openpyxl import load_workbook
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from app.models.document import Document
from app.services.anonymization_service import (
    build_replacement_map,
    remove_overlapping_detections,
)
from app.services.file_service import (
    decrypt_and_verify_file,
)
from app.services.pii_detector import detect_pii
from app.services.text_extractor import extract_text


ANONYMIZED_STORAGE = (
    Path("storage")
    / "anonymized"
)

ANONYMIZED_STORAGE.mkdir(
    parents=True,
    exist_ok=True
)


def build_string_replacement_map(
    replacement_map: dict[
        tuple[str, str],
        str
    ]
) -> dict[str, str]:

    string_map: dict[str, str] = {}

    for (
        _pii_type,
        original_value
    ), replacement in replacement_map.items():

        if original_value not in string_map:
            string_map[
                original_value
            ] = replacement

    return string_map


def replace_text(
    text: str,
    replacement_map: dict[str, str]
) -> str:

    result = text

    ordered_replacements = sorted(
        replacement_map.items(),
        key=lambda item: len(
            item[0]
        ),
        reverse=True,
    )

    for (
        original_value,
        replacement
    ) in ordered_replacements:

        result = result.replace(
            original_value,
            replacement
        )

    return result


def prepare_anonymization(
    data: bytes,
    file_type: str
) -> tuple[
    dict[str, str],
    list[dict]
]:

    extracted_text = extract_text(
        data=data,
        file_type=file_type,
    )

    detections = detect_pii(
        extracted_text
    )

    detections = (
        remove_overlapping_detections(
            detections
        )
    )

    typed_replacement_map = (
        build_replacement_map(
            detections
        )
    )

    string_replacement_map = (
        build_string_replacement_map(
            typed_replacement_map
        )
    )

    return (
        string_replacement_map,
        detections,
    )


def anonymize_txt(
    data: bytes,
    replacement_map: dict[str, str]
) -> bytes:

    text = data.decode(
        "utf-8"
    )

    anonymized_text = replace_text(
        text,
        replacement_map
    )

    return anonymized_text.encode(
        "utf-8"
    )


def anonymize_csv(
    data: bytes,
    replacement_map: dict[str, str]
) -> bytes:

    text = data.decode(
        "utf-8"
    )

    input_stream = io.StringIO(
        text
    )

    try:
        dialect = csv.Sniffer().sniff(
            text[:2048]
        )

    except csv.Error:
        dialect = csv.excel

    reader = csv.reader(
        input_stream,
        dialect
    )

    output_stream = io.StringIO()

    writer = csv.writer(
        output_stream,
        dialect
    )

    for row in reader:

        anonymized_row = [
            replace_text(
                cell,
                replacement_map
            )
            for cell in row
        ]

        writer.writerow(
            anonymized_row
        )

    return output_stream.getvalue().encode(
        "utf-8"
    )


def anonymize_docx(
    data: bytes,
    replacement_map: dict[str, str]
) -> bytes:

    document = DocxDocument(
        io.BytesIO(data)
    )

    for paragraph in document.paragraphs:

        paragraph.text = replace_text(
            paragraph.text,
            replacement_map
        )

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    paragraph.text = (
                        replace_text(
                            paragraph.text,
                            replacement_map
                        )
                    )

    output_stream = io.BytesIO()

    document.save(
        output_stream
    )

    return output_stream.getvalue()


def anonymize_xlsx(
    data: bytes,
    replacement_map: dict[str, str]
) -> bytes:

    workbook = load_workbook(
        filename=io.BytesIO(data),
        data_only=False,
    )

    try:

        for worksheet in workbook.worksheets:

            for row in worksheet.iter_rows():

                for cell in row:

                    if not isinstance(
                        cell.value,
                        str
                    ):
                        continue

                    if cell.value.startswith("="):
                        continue

                    cell.value = replace_text(
                        cell.value,
                        replacement_map
                    )

        output_stream = io.BytesIO()

        workbook.save(
            output_stream
        )

        return output_stream.getvalue()

    finally:
        workbook.close()


def anonymize_pdf(
    data: bytes,
    replacement_map: dict[str, str]
) -> bytes:

    extracted_text = extract_text(
        data=data,
        file_type="pdf",
    )

    anonymized_text = replace_text(
        extracted_text,
        replacement_map
    )

    output_stream = io.BytesIO()

    document = SimpleDocTemplate(
        output_stream
    )

    styles = getSampleStyleSheet()

    story = []

    for line in anonymized_text.splitlines():

        if line.strip():

            safe_line = escape(
                line
            )

            story.append(
                Paragraph(
                    safe_line,
                    styles["BodyText"]
                )
            )

        else:

            story.append(
                Spacer(
                    1,
                    8
                )
            )

    document.build(
        story
    )

    return output_stream.getvalue()


def anonymize_document_bytes(
    data: bytes,
    file_type: str
) -> dict:

    normalized_file_type = (
        file_type
        .lower()
        .strip()
        .lstrip(".")
    )

    replacement_map, detections = (
        prepare_anonymization(
            data=data,
            file_type=normalized_file_type,
        )
    )

    anonymizers = {
        "txt": anonymize_txt,
        "csv": anonymize_csv,
        "docx": anonymize_docx,
        "xlsx": anonymize_xlsx,
        "pdf": anonymize_pdf,
    }

    anonymizer = anonymizers.get(
        normalized_file_type
    )

    if anonymizer is None:
        raise ValueError(
            "Unsupported file type "
            "for anonymization"
        )

    output_data = anonymizer(
        data,
        replacement_map
    )

    pii_counts: dict[str, int] = {}

    for detection in detections:

        pii_type = detection[
            "type"
        ]

        pii_counts[pii_type] = (
            pii_counts.get(
                pii_type,
                0
            )
            + 1
        )

    return {
        "data":
            output_data,

        "pii_count":
            len(detections),

        "pii_counts":
            pii_counts,

        "replacement_count":
            len(replacement_map),
    }


def anonymize_and_store_document(
    document: Document
) -> dict:

    decrypted_data = (
        decrypt_and_verify_file(
            encrypted_file_path=(
                document.encrypted_file_path
            ),
            expected_sha256=(
                document.sha256_hash
            ),
        )
    )

    result = anonymize_document_bytes(
        data=decrypted_data,
        file_type=document.file_type,
    )

    extension = (
        document.file_type
        .lower()
        .strip()
        .lstrip(".")
    )

    output_filename = (
        f"{uuid.uuid4().hex}."
        f"{extension}"
    )

    output_path = (
        ANONYMIZED_STORAGE
        / output_filename
    )

    output_path.write_bytes(
        result["data"]
    )

    return {
        "anonymized_file_path":
            str(output_path),

        "pii_count":
            result["pii_count"],

        "pii_counts":
            result["pii_counts"],

        "replacement_count":
            result["replacement_count"],
    }