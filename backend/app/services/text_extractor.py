import csv
import io

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader


SUPPORTED_FILE_TYPES = {
    "pdf",
    "docx",
    "txt",
    "xlsx",
    "csv",
}


def extract_pdf_text(
    data: bytes
) -> str:

    pdf_stream = io.BytesIO(
        data
    )

    reader = PdfReader(
        pdf_stream
    )

    pages: list[str] = []

    for page in reader.pages:
        page_text = (
            page.extract_text()
            or ""
        )

        if page_text.strip():
            pages.append(
                page_text
            )

    return "\n".join(
        pages
    )


def extract_docx_text(
    data: bytes
) -> str:

    document = DocxDocument(
        io.BytesIO(data)
    )

    content: list[str] = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            content.append(
                text
            )

    for table in document.tables:

        for row in table.rows:

            values = [
                cell.text.strip()
                for cell in row.cells
            ]

            if any(values):
                content.append(
                    " | ".join(values)
                )

    return "\n".join(
        content
    )


def extract_txt_text(
    data: bytes
) -> str:

    return data.decode(
        "utf-8"
    )


def extract_csv_text(
    data: bytes
) -> str:

    text = data.decode(
        "utf-8"
    )

    reader = csv.reader(
        io.StringIO(text)
    )

    rows: list[str] = []

    for row in reader:

        values = [
            str(value).strip()
            for value in row
        ]

        rows.append(
            " | ".join(values)
        )

    return "\n".join(
        rows
    )


def extract_xlsx_text(
    data: bytes
) -> str:

    workbook = load_workbook(
        filename=io.BytesIO(data),
        read_only=True,
        data_only=True,
    )

    content: list[str] = []

    try:

        for worksheet in workbook.worksheets:

            content.append(
                f"Sheet: {worksheet.title}"
            )

            for row in worksheet.iter_rows(
                values_only=True
            ):

                values = [
                    ""
                    if value is None
                    else str(value)
                    for value in row
                ]

                if any(
                    value.strip()
                    for value in values
                ):
                    content.append(
                        " | ".join(values)
                    )

    finally:
        workbook.close()

    return "\n".join(
        content
    )


def extract_text(
    data: bytes,
    file_type: str
) -> str:

    normalized_file_type = (
        file_type
        .lower()
        .strip()
        .lstrip(".")
    )

    extractors = {
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "txt": extract_txt_text,
        "xlsx": extract_xlsx_text,
        "csv": extract_csv_text,
    }

    extractor = extractors.get(
        normalized_file_type
    )

    if extractor is None:
        raise ValueError(
            "Unsupported file type "
            "for text extraction"
        )

    return extractor(
        data
    )